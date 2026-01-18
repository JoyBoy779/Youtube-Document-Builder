from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import TranscriptsDisabled
import re

user_input = input("Enter YouTube video URL or ID: ").strip()
# If input looks like a URL, extract video ID
match = re.search(r"(?:v=|youtu\.be/)([A-Za-z0-9_-]{11})", user_input) # video id is always 11 charaters 
if match:
    video_id = match.group(1)
else:
    # Assume user directly entered video ID
    video_id = user_input

# Get language code from user (e.g. en, hi)
lang = input("Enter transcript language code (e.g. en, hi): ").strip().lower()
try:
    # Fetch transcript (Hindi)
    fetched_transcript = YouTubeTranscriptApi().fetch(
        video_id,
        languages=[lang]
    )
    # Convert transcript chunks into plain text
    transcript_list = fetched_transcript.to_raw_data()
    transcript = " ".join(chunk["text"] for chunk in transcript_list)

except TranscriptsDisabled:
    raise RuntimeError("No captions available for this video.")

except Exception as e:
    raise RuntimeError(f"Invalid video ID or transcript fetch failed: {e}")

#storing in a doc
from docx import Document
doc = Document()
doc.add_heading("YouTube Transcript", level=1)
doc.add_paragraph(transcript)

file_name = f"{lang}_transcript_{video_id}.docx"
doc.save(file_name)
print(f"Transcript saved as {file_name}")

